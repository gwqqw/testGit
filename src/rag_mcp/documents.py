from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import hashlib


SUPPORTED_SUFFIXES = {".txt", ".md", ".docx"}


@dataclass(frozen=True)
class Document:
    doc_id: str
    path: Path
    content: str


def doc_id_for(path: Path, root: Path) -> str:
    rel_path = path.relative_to(root).as_posix()
    digest = hashlib.sha256(rel_path.encode("utf-8")).hexdigest()
    return digest


def _read_text_file(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def _read_docx_file(path: Path) -> str:
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:
        raise RuntimeError("python-docx is required to read .docx files") from exc

    doc = DocxDocument(path)
    lines = [paragraph.text for paragraph in doc.paragraphs if paragraph.text]
    return "\n".join(lines)


def read_document(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md"}:
        return _read_text_file(path)
    if suffix == ".docx":
        return _read_docx_file(path)
    raise ValueError(f"Unsupported file type: {suffix}")


def load_documents(docs_dir: Path) -> list[Document]:
    if not docs_dir.exists():
        return []

    documents: list[Document] = []
    for path in sorted(docs_dir.rglob("*")):
        if not path.is_file():
            continue
        if path.suffix.lower() not in SUPPORTED_SUFFIXES:
            continue
        content = read_document(path)
        if not content.strip():
            continue
        doc_id = doc_id_for(path, docs_dir)
        documents.append(Document(doc_id=doc_id, path=path, content=content))
    return documents